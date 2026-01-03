# converter/utils.py
import os
import io
import tempfile
import zipfile  # <-- ADD THIS IMPORT
import PyPDF2
from pdf2docx import Converter
from docx import Document
from PIL import Image
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4, A5, legal
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.pdfgen import canvas
from django.core.files.base import ContentFile
from django.utils import timezone

def convert_pdf_to_word(pdf_path, output_format='docx', preserve_layout=True, 
                       use_ocr=False, extract_text_only=False):
    """
    Convert PDF to Word with formatting options
    
    Args:
        pdf_path: Path to PDF file
        output_format: 'docx', 'doc', 'rtf', or 'txt'
        preserve_layout: Keep original layout
        use_ocr: Use OCR for scanned documents
        extract_text_only: Extract only text, no images/tables
    """
    try:
        # For DOCX format (recommended)
        if output_format == 'docx':
            # Create a temporary file for the docx output
            docx_stream = io.BytesIO()
            
            # Convert PDF to DOCX
            cv = Converter(pdf_path)
            
            if preserve_layout:
                # Try to preserve layout
                cv.convert(docx_stream, start=0, end=None)
            else:
                # Simple conversion
                cv.convert(docx_stream)
            
            cv.close()
            docx_stream.seek(0)
            return docx_stream
        
        # For TXT format (text only)
        elif output_format == 'txt':
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            
            # Return as BytesIO
            return io.BytesIO(text.encode('utf-8'))
        
        # For DOC format (older Word format)
        elif output_format == 'doc':
            # Convert to DOCX first
            docx_stream = io.BytesIO()
            cv = Converter(pdf_path)
            cv.convert(docx_stream)
            cv.close()
            docx_stream.seek(0)
            
            # For now, return as DOCX since .doc is tricky
            return docx_stream
        
        # For RTF format
        elif output_format == 'rtf':
            # Convert PDF to text first, then format as RTF
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            
            # Create simple RTF header and content
            rtf_content = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}\f0\fs24 "
            rtf_content += text.replace('\n', '\\par ')
            rtf_content += "}"
            
            return io.BytesIO(rtf_content.encode('utf-8'))
        
        # Default to DOCX
        else:
            docx_stream = io.BytesIO()
            cv = Converter(pdf_path)
            cv.convert(docx_stream)
            cv.close()
            docx_stream.seek(0)
            return docx_stream
            
    except Exception as e:
        # Fallback: Extract text only
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            
            return io.BytesIO(text.encode('utf-8'))
        except:
            raise Exception(f"Conversion error: {str(e)}")

def convert_word_to_pdf(word_path):
    """Convert Word document to PDF - Linux compatible version."""
    try:
        # Try using python-docx and reportlab for Linux compatibility
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        
        # Read Word document
        doc = Document(word_path)
        
        # Create PDF buffer
        pdf_buffer = io.BytesIO()
        
        # Create PDF document
        pdf_doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Extract content from Word
        story = []
        styles = getSampleStyleSheet()
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                # Preserve some formatting
                text = paragraph.text
                
                # Check for heading styles
                if paragraph.style.name.startswith('Heading'):
                    style = styles['Heading1'] if 'Heading1' in paragraph.style.name else styles['Heading2']
                else:
                    style = styles['Normal']
                
                story.append(Paragraph(text, style))
                story.append(Spacer(1, 12))
        
        # Process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            
            if table_data:
                from reportlab.platypus import Table, TableStyle
                from reportlab.lib import colors
                
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 20))
        
        # Build PDF
        if story:
            pdf_doc.build(story)
        else:
            # Fallback: create simple PDF with text
            c = canvas.Canvas(pdf_buffer, pagesize=letter)
            y = 750
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if y < 50:
                        c.showPage()
                        y = 750
                    
                    c.drawString(50, y, paragraph.text[:100])
                    y -= 20
            
            c.save()
        
        pdf_buffer.seek(0)
        return pdf_buffer
        
    except Exception as e:
        # Fallback: create a very basic PDF
        try:
            pdf_buffer = io.BytesIO()
            c = canvas.Canvas(pdf_buffer, pagesize=letter)
            
            with open(word_path, 'rb') as f:
                # Read as binary and extract text if possible
                import docx2txt
                text = docx2txt.process(word_path)
                
                y = 750
                lines = text.split('\n')
                
                for line in lines:
                    if y < 50:
                        c.showPage()
                        y = 750
                    
                    c.drawString(50, y, line[:100])
                    y -= 20
            
            c.save()
            pdf_buffer.seek(0)
            return pdf_buffer
            
        except Exception as fallback_error:
            raise Exception(f"Word to PDF conversion failed: {str(fallback_error)}")

def merge_pdfs(pdf_paths):
    """Merge multiple PDFs into one - Render compatible."""
    try:
        merger = PyPDF2.PdfMerger()
        
        for pdf_path in pdf_paths:
            # Check if file exists
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"PDF file not found: {pdf_path}")
            
            # Read file content first to avoid path issues
            with open(pdf_path, 'rb') as file:
                pdf_data = io.BytesIO(file.read())
                pdf_data.seek(0)
                merger.append(pdf_data)
        
        output_stream = io.BytesIO()
        merger.write(output_stream)
        merger.close()
        
        output_stream.seek(0)
        return output_stream
        
    except Exception as e:
        raise Exception(f"PDF merge failed: {str(e)}")

def split_pdf(pdf_path, pages):
    """Split PDF by pages."""
    try:
        pages_to_extract = []
        
        # Handle different page formats
        if '-' in pages:
            # Range format
            start, end = map(int, pages.split('-'))
            pages_to_extract = list(range(start-1, end))  # 0-indexed
        elif ',' in pages:
            # List format
            parts = pages.split(',')
            for part in parts:
                part = part.strip()
                if '-' in part:
                    start, end = map(int, part.split('-'))
                    pages_to_extract.extend(range(start-1, end))
                else:
                    pages_to_extract.append(int(part) - 1)
        else:
            # Single page
            pages_to_extract = [int(pages) - 1]
        
        # Extract pages
        pdf_reader = PyPDF2.PdfReader(pdf_path)
        pdf_writer = PyPDF2.PdfWriter()
        
        for page_num in pages_to_extract:
            if 0 <= page_num < len(pdf_reader.pages):
                pdf_writer.add_page(pdf_reader.pages[page_num])
        
        output_stream = io.BytesIO()
        pdf_writer.write(output_stream)
        output_stream.seek(0)
        
        return output_stream
        
    except Exception as e:
        raise Exception(f"PDF split failed: {str(e)}")

def split_pdf_by_range(pdf_path, pages):
    """Split PDF by specific page ranges."""
    # Parse page ranges
    pages_to_extract = parse_page_ranges(pages)
    
    # Read PDF
    pdf_reader = PyPDF2.PdfReader(pdf_path)
    total_pages = len(pdf_reader.pages)
    
    # Create zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for i, page_range in enumerate(pages_to_extract):
            pdf_writer = PyPDF2.PdfWriter()
            
            # Adjust for 0-indexed pages
            start = max(0, page_range['start'] - 1)
            end = min(total_pages, page_range['end'])
            
            for page_num in range(start, end):
                pdf_writer.add_page(pdf_reader.pages[page_num])
            
            # Create PDF in memory
            pdf_buffer = io.BytesIO()
            pdf_writer.write(pdf_buffer)
            pdf_buffer.seek(0)
            
            filename = f"part_{i+1}_pages_{page_range['start']}-{page_range['end']}.pdf"
            zip_file.writestr(filename, pdf_buffer.read())
    
    zip_buffer.seek(0)
    return zip_buffer

def split_pdf_every_page(pdf_path, split_every=1):
    """Split PDF into files with specified number of pages each."""
    pdf_reader = PyPDF2.PdfReader(pdf_path)
    total_pages = len(pdf_reader.pages)
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        part_num = 1
        for start in range(0, total_pages, split_every):
            pdf_writer = PyPDF2.PdfWriter()
            end = min(start + split_every, total_pages)
            
            for page_num in range(start, end):
                pdf_writer.add_page(pdf_reader.pages[page_num])
            
            pdf_buffer = io.BytesIO()
            pdf_writer.write(pdf_buffer)
            pdf_buffer.seek(0)
            
            filename = f"part_{part_num}_pages_{start+1}-{end}.pdf"
            zip_file.writestr(filename, pdf_buffer.read())
            part_num += 1
    
    zip_buffer.seek(0)
    return zip_buffer

def split_pdf_by_count(pdf_path, pages_per_file=10):
    """Split PDF by number of pages per file."""
    return split_pdf_every_page(pdf_path, pages_per_file)

def split_pdf_custom(pdf_path, split_points):
    """Split PDF at custom split points."""
    # Parse split points
    points = [int(p.strip()) for p in split_points.split(',') if p.strip().isdigit()]
    points = sorted(set(points))
    
    pdf_reader = PyPDF2.PdfReader(pdf_path)
    total_pages = len(pdf_reader.pages)
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        start = 0
        part_num = 1
        
        for split_point in points:
            if split_point > start and split_point <= total_pages:
                pdf_writer = PyPDF2.PdfWriter()
                for page_num in range(start, split_point):
                    pdf_writer.add_page(pdf_reader.pages[page_num])
                
                pdf_buffer = io.BytesIO()
                pdf_writer.write(pdf_buffer)
                pdf_buffer.seek(0)
                
                filename = f"part_{part_num}_pages_{start+1}-{split_point}.pdf"
                zip_file.writestr(filename, pdf_buffer.read())
                
                start = split_point
                part_num += 1
        
        # Add remaining pages
        if start < total_pages:
            pdf_writer = PyPDF2.PdfWriter()
            for page_num in range(start, total_pages):
                pdf_writer.add_page(pdf_reader.pages[page_num])
            
            pdf_buffer = io.BytesIO()
            pdf_writer.write(pdf_buffer)
            pdf_buffer.seek(0)
            
            filename = f"part_{part_num}_pages_{start+1}-{total_pages}.pdf"
            zip_file.writestr(filename, pdf_buffer.read())
    
    zip_buffer.seek(0)
    return zip_buffer

def parse_page_ranges(pages_str):
    """Parse page range string like '1-3, 5, 7-10'."""
    ranges = []
    parts = pages_str.split(',')
    
    for part in parts:
        part = part.strip()
        if '-' in part:
            start_str, end_str = part.split('-')
            start = int(start_str.strip())
            end = int(end_str.strip())
            ranges.append({'start': start, 'end': end})
        elif part.isdigit():
            page = int(part)
            ranges.append({'start': page, 'end': page})
    
    return ranges

# Also update the existing split_pdf function to use the new one
def split_pdf(pdf_path, pages):
    """Legacy function - splits PDF by page ranges."""
    return split_pdf_by_range(pdf_path, pages)

def compress_pdf(input_path, compression_level='medium', optimize_images=True, 
                 optimize_fonts=False, remove_metadata=False):
    """
    Compress a PDF file. This is a placeholder implementation.
    In a real application, you would use a PDF compression library.
    
    Args:
        input_path: Path to the input PDF file
        compression_level: 'low', 'medium', or 'high'
        optimize_images: Whether to compress images
        optimize_fonts: Whether to optimize fonts
        remove_metadata: Whether to remove metadata
        
    Returns:
        ContentFile: Compressed PDF as a Django ContentFile
    """
    try:
        # Make sure input file exists
        if not os.path.exists(input_path):
            raise ValueError(f"Input file does not exist: {input_path}")
        
        # Read the input file
        with open(input_path, 'rb') as f:
            pdf_data = f.read()
        
        # In a real implementation, you would:
        # 1. Use PyPDF2, pikepdf, or PyMuPDF to compress the PDF
        # 2. Apply compression based on the parameters
        # 3. Return the compressed PDF
        
        # For now, we'll return the original file as a placeholder
        # TODO: Implement actual PDF compression
        
        # Create a ContentFile with the PDF data
        return ContentFile(pdf_data, name='compressed.pdf')
        
    except Exception as e:
        raise Exception(f"Error compressing PDF: {str(e)}")

def compress_pdf_with_pypdf2(input_path, compression_level='medium', 
                           optimize_images=True, optimize_fonts=False, 
                           remove_metadata=False):
    """
    Alternative implementation using PyPDF2.
    Note: PyPDF2 has limited compression capabilities.
    """
    try:
        import PyPDF2
        
        # Create output in memory
        output = ContentFile(b'', name='compressed.pdf')
        
        # Read input PDF
        with open(input_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            writer = PyPDF2.PdfWriter()
            
            # Copy all pages
            for page in reader.pages:
                writer.add_page(page)
            
            # Apply compression settings
            if compression_level == 'high':
                # High compression - may reduce quality
                writer.compress_content_streams = True
            elif compression_level == 'medium':
                # Medium compression
                writer.compress_content_streams = True
            # Low compression keeps as-is
            
            # Write to output
            writer.write(output)
            output.seek(0)
            
        return output
        
    except ImportError:
        # PyPDF2 not installed, fall back to simple method
        return compress_pdf(input_path, compression_level, optimize_images, 
                          optimize_fonts, remove_metadata)
    except Exception as e:
        raise Exception(f"Error with PyPDF2 compression: {str(e)}")

def compress_pdf_with_pikepdf(input_path, compression_level='medium',
                            optimize_images=True, optimize_fonts=False,
                            remove_metadata=False):
    """
    Better compression using pikepdf library.
    """
    try:
        import pikepdf
        
        # Open the PDF
        pdf = pikepdf.open(input_path)
        
        # Apply compression based on level
        if compression_level != 'low':
            # Save with compression (pikepdf automatically applies some compression)
            output = ContentFile(b'', name='compressed.pdf')
            pdf.save(output, compress_streams=True)
            output.seek(0)
            
            # Additional optimizations
            if optimize_images:
                # pikepdf doesn't have direct image optimization
                # You might need additional libraries for this
                pass
                
            if remove_metadata:
                # Remove metadata
                with pikepdf.open(output) as pdf_without_metadata:
                    if '/Metadata' in pdf_without_metadata.Root:
                        del pdf_without_metadata.Root.Metadata
                    pdf_without_metadata.save(output, compress_streams=True)
                    output.seek(0)
                    
            return output
        else:
            # For low compression, just return a copy
            with open(input_path, 'rb') as f:
                return ContentFile(f.read(), name='compressed.pdf')
                
    except ImportError:
        # pikepdf not installed, fall back to PyPDF2
        return compress_pdf_with_pypdf2(input_path, compression_level,
                                      optimize_images, optimize_fonts,
                                      remove_metadata)
    except Exception as e:
        raise Exception(f"Error with pikepdf compression: {str(e)}")

def compress_pdf_advanced(pdf_path, compression_level='medium', 
                         optimize_images=True, remove_metadata=False, 
                         downsample_images=True):
    """
    Enhanced PDF compression with multiple optimization options
    
    Args:
        pdf_path: Path to PDF file
        compression_level: 'low', 'medium', or 'high'
        optimize_images: Reduce image quality
        remove_metadata: Remove document metadata
        downsample_images: Reduce image resolution
    """
    try:
        import PyPDF2
        import io
        
        pdf_reader = PyPDF2.PdfReader(pdf_path)
        pdf_writer = PyPDF2.PdfWriter()
        
        # Copy all pages with optimizations
        for page_num, page in enumerate(pdf_reader.pages):
            pdf_writer.add_page(page)
        
        # Apply compression based on level
        if compression_level == 'high':
            # Maximum compression
            for page in pdf_writer.pages:
                page.compress_content_streams()
                # Additional high compression techniques
                if hasattr(page, 'images'):
                    for image in page.images:
                        # Reduce image quality
                        image.compress()
        
        elif compression_level == 'medium':
            # Medium compression (default)
            for page in pdf_writer.pages:
                page.compress_content_streams()
        
        # Apply additional optimizations if requested
        if optimize_images:
            for page in pdf_writer.pages:
                if hasattr(page, 'images'):
                    for image in page.images:
                        image.compress()
        
        if remove_metadata:
            # Remove metadata
            if '/Metadata' in pdf_writer._root_object:
                del pdf_writer._root_object['/Metadata']
        
        if downsample_images:
            # Downsample images (simplified - in production use proper image processing)
            for page in pdf_writer.pages:
                if hasattr(page, 'images'):
                    for image in page.images:
                        # This is a placeholder - actual downsampling would require PIL
                        pass
        
        # Write to output stream
        output_stream = io.BytesIO()
        pdf_writer.write(output_stream)
        output_stream.seek(0)
        
        return output_stream
        
    except Exception as e:
        # Fallback to basic compression
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_path)
            pdf_writer = PyPDF2.PdfWriter()
            
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)
            
            if compression_level == 'high':
                for page in pdf_writer.pages:
                    page.compress_content_streams()
            
            output_stream = io.BytesIO()
            pdf_writer.write(output_stream)
            output_stream.seek(0)
            return output_stream
            
        except Exception as fallback_error:
            raise Exception(f"PDF compression failed: {str(fallback_error)}")

def convert_excel_to_pdf(excel_path, include_gridlines=True, fit_to_page=True, include_headers=True):
    """Convert Excel to PDF with options."""
    try:
        # Import reportlab components at the top of the function
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.lib.units import inch
        
        # Read Excel file
        df = pd.read_excel(excel_path)
        
        # Create PDF buffer
        pdf_buffer = io.BytesIO()
        
        # Configure page size based on fit_to_page option
        if fit_to_page:
            # Calculate required width and height
            num_cols = len(df.columns)
            num_rows = len(df) + 1  # +1 for header row
            
            # Approximate cell sizes
            col_width = 1.5 * inch
            row_height = 0.4 * inch
            
            required_width = num_cols * col_width
            required_height = num_rows * row_height
            
            # Use letter size or adjust if content is too large
            if required_width > 10*inch or required_height > 10*inch:
                pagesize = (required_width + 2*inch, required_height + 2*inch)
            else:
                pagesize = letter
        else:
            pagesize = letter
        
        # Create SimpleDocTemplate with the pagesize
        doc = SimpleDocTemplate(pdf_buffer, pagesize=pagesize)
        
        # Convert DataFrame to list of lists for ReportLab
        data = [df.columns.tolist()] + df.values.tolist()
        
        # Create table with appropriate column widths
        col_widths = [doc.width/len(df.columns)] * len(df.columns)
        table = Table(data, colWidths=col_widths)
        
        # Add style based on options
        style_commands = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4CAF50')),  # Green header
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F9F9F9')),
        ]
        
        if include_gridlines:
            style_commands.append(('GRID', (0, 0), (-1, -1), 0.5, colors.grey))
        else:
            style_commands.append(('BOX', (0, 0), (-1, -1), 1, colors.black))
        
        style = TableStyle(style_commands)
        table.setStyle(style)
        
        # Build PDF with optional headers
        elements = []
        
        if include_headers:
            # Add title
            from reportlab.platypus import Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            styles = getSampleStyleSheet()
            title = Paragraph("Excel to PDF Conversion", styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 20))
        
        elements.append(table)
        doc.build(elements)
        pdf_buffer.seek(0)
        
        return pdf_buffer
        
    except Exception as e:
        raise Exception(f"Excel to PDF conversion failed: {str(e)}")

def convert_images_to_pdf(image_paths, page_size='A4', orientation='portrait', 
                         placement='fit', add_page_numbers=False):
    """Convert images to PDF with various layout options."""
    try:
        # Map page sizes
        size_map = {
            'A4': A4,
            'letter': letter,
            'A5': A5,
            'legal': legal,
            'A4': A4  # default
        }
        
        page_size_obj = size_map.get(page_size, A4)
        
        # Adjust for orientation
        if orientation == 'landscape':
            page_size_obj = (page_size_obj[1], page_size_obj[0])
        
        # Create PDF buffer
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=page_size_obj)
        
        # Set up page number variables
        page_number = 1
        total_pages = len(image_paths)
        
        for idx, img_path in enumerate(image_paths):
            try:
                # Open image
                img = Image.open(img_path)
                img_width, img_height = img.size
                page_width, page_height = page_size_obj
                
                # Calculate scaling based on placement option
                if placement == 'fit':
                    # Fit to page with margins
                    width_ratio = (page_width * 0.9) / img_width
                    height_ratio = (page_height * 0.9) / img_height
                    scale = min(width_ratio, height_ratio)
                    scaled_width = img_width * scale
                    scaled_height = img_height * scale
                    x = (page_width - scaled_width) / 2
                    y = (page_height - scaled_height) / 2
                    
                elif placement == 'full':
                    # Full page (stretch to fill)
                    scaled_width = page_width
                    scaled_height = page_height
                    x = 0
                    y = 0
                    
                elif placement == 'center':
                    # Center original size
                    scaled_width = img_width
                    scaled_height = img_height
                    x = (page_width - scaled_width) / 2
                    y = (page_height - scaled_height) / 2
                    
                else:  # 'multiple' or default to fit
                    # For multiple images per page (simplified - one per page for now)
                    width_ratio = page_width / img_width
                    height_ratio = page_height / img_height
                    scale = min(width_ratio, height_ratio) * 0.8
                    scaled_width = img_width * scale
                    scaled_height = img_height * scale
                    x = (page_width - scaled_width) / 2
                    y = (page_height - scaled_height) / 2
                
                # Draw image on PDF
                c.drawImage(img_path, x, y, width=scaled_width, height=scaled_height)
                
                # Add page numbers if requested
                if add_page_numbers:
                    c.setFont("Helvetica", 10)
                    c.setFillColorRGB(0.5, 0.5, 0.5)  # Gray color
                    c.drawRightString(page_width - 50, 30, f"Page {page_number} of {total_pages}")
                    page_number += 1
                
                # Add image info (optional)
                c.setFont("Helvetica", 8)
                c.setFillColorRGB(0.3, 0.3, 0.3)
                img_name = os.path.basename(img_path)
                c.drawString(50, 30, f"Image: {img_name}")
                
                c.showPage()  # New page for next image
                
            except Exception as img_error:
                print(f"Error processing image {img_path}: {img_error}")
                # Still create a page with error message
                c.setFont("Helvetica", 12)
                c.drawString(100, page_height/2, f"Error loading image: {os.path.basename(img_path)}")
                c.showPage()
                continue
        
        c.save()
        pdf_buffer.seek(0)
        
        return pdf_buffer
        
    except Exception as e:
        raise Exception(f"Image to PDF conversion failed: {str(e)}")

def handle_file_upload(file, request):
    """Handle file upload and create record."""
    from .models import UploadedFile
    
    uploaded = UploadedFile.objects.create(
        original_filename=file.name,
        file_type=os.path.splitext(file.name)[1].lower(),
        session_key=request.session.session_key or 'anonymous'
    )
    uploaded.file.save(file.name, file)
    return uploaded


# In your utils.py or context_processors.py, update the get_all_tools function:
def get_all_tools():
    """Return a list of all available conversion tools."""
    tools = [
        {
            'name': 'PDF to Word',
            'description': 'Convert PDF files to editable Word documents',
            'url': 'pdf_to_word',  # Changed from 'pdf-to-word' to 'pdf_to_word'
            'icon': 'fa-file-word',
            'color': 'blue'
        },
        {
            'name': 'Word to PDF',
            'description': 'Convert Word documents to PDF format',
            'url': 'word_to_pdf',  # Changed from 'word-to-pdf' to 'word_to_pdf'
            'icon': 'fa-file-pdf',
            'color': 'red'
        },
        {
            'name': 'Merge PDF',
            'description': 'Combine multiple PDF files into one',
            'url': 'merge_pdf',  # Changed from 'merge-pdf' to 'merge_pdf'
            'icon': 'fa-copy',
            'color': 'purple'
        },
        {
            'name': 'Split PDF',
            'description': 'Split PDF into multiple files or extract pages',
            'url': 'split_pdf',  # Changed from 'split-pdf' to 'split_pdf'
            'icon': 'fa-cut',
            'color': 'green'
        },
        {
            'name': 'Compress PDF',
            'description': 'Reduce PDF file size without losing quality',
            'url': 'compress_pdf',  # Changed from 'compress-pdf' to 'compress_pdf'
            'icon': 'fa-compress',
            'color': 'yellow'
        },
        {
            'name': 'Excel to PDF',
            'description': 'Convert Excel spreadsheets to PDF',
            'url': 'excel_to_pdf',  # Changed from 'excel-to-pdf' to 'excel_to_pdf'
            'icon': 'fa-file-excel',
            'color': 'green'
        },
        {
            'name': 'Image to PDF',
            'description': 'Convert images to PDF documents',
            'url': 'image_to_pdf',  # Changed from 'image-to-pdf' to 'image_to_pdf'
            'icon': 'fa-image',
            'color': 'pink'
        },
    ]
    return tools